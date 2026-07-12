"""Runnable, DB/network-free check for IngestProcessor._extract_text's format
branches (txt/html/pdf already existed; docx/image are new). Builds tiny fixtures
in-memory with python-docx + a hand-built PNG, and stubs describe_image so the
image branch is exercised without a live VLM call.

For the full Phase A/B pipeline against real Postgres/Stratus, see self_check.py.

Run: python catalyst_functions/ingest_processor/test_extract_formats.py
"""
from __future__ import annotations

import struct
import sys
import zlib
from io import BytesIO
from pathlib import Path
from unittest import mock

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from catalyst_functions.ingest_processor.pipeline import processor as processor_module  # noqa: E402
from catalyst_functions.ingest_processor.pipeline.processor import IngestProcessor  # noqa: E402


def _make_test_png() -> bytes:
    def chunk(tag: bytes, payload: bytes) -> bytes:
        return struct.pack(">I", len(payload)) + tag + payload + struct.pack(">I", zlib.crc32(tag + payload))

    width, height = 4, 4
    raw = (bytes([0]) + bytes((10, 20, 30)) * width) * height
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")


def main() -> int:
    extractor = IngestProcessor.__new__(IngestProcessor)  # skips __init__ (no DB connection needed)

    assert extractor._extract_text("f.txt", b"hello world") == "hello world"
    print("[test_extract_formats] PASS: .txt")

    html = b"<html><body><p>Case <b>42</b></p></body></html>"
    assert "Case" in extractor._extract_text("f.html", html) and "42" in extractor._extract_text("f.html", html)
    print("[test_extract_formats] PASS: .html")

    doc = Document()
    doc.add_paragraph("Complainant: Ravi Kumar")
    doc.add_paragraph("Account: 1234567890")
    buf = BytesIO()
    doc.save(buf)
    docx_text = extractor._extract_text("fir.docx", buf.getvalue())
    assert "Ravi Kumar" in docx_text and "1234567890" in docx_text, docx_text
    print("[test_extract_formats] PASS: .docx (python-docx)")

    with mock.patch.object(processor_module, "describe_image", return_value="Account: 9998887770") as mocked:
        image_text = extractor._extract_text("screenshot.png", _make_test_png())
    mocked.assert_called_once()
    assert image_text == "Account: 9998887770"
    print("[test_extract_formats] PASS: .png dispatches to describe_image (VLM)")

    print("\n[test_extract_formats] ALL PASS")
    return 0


if __name__ == "__main__":
    sys.exit(main())
